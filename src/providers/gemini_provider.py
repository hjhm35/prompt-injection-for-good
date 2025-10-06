from typing import List
import os
import mimetypes
from pathlib import Path
import time
from datetime import datetime, timedelta

try:
    import google.generativeai as genai
    from google.generativeai.types import HarmCategory, HarmBlockThreshold
except ImportError:
    raise ImportError("Google Generative AI library not installed. Run: pip install google-generativeai")

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

from .base import BaseProvider, EvaluationResult

class GeminiProvider(BaseProvider):
    """Google Gemini provider implementation"""
    
    # No class-level caching - app-level cache handles this
    
    # Gemini supports extensive file formats
    SUPPORTED_FILE_TYPES = {
        # Documents
        'application/pdf': '.pdf',
        # Note: DOCX files require text extraction - not supported for direct upload
        'application/msword': '.doc',
        'text/plain': '.txt',
        'text/markdown': '.md',
        'text/csv': '.csv',
        'message/rfc822': '.eml',
        'text/email': '.eml',
        'application/json': '.json',
        # Images
        'image/jpeg': '.jpg',
        'image/png': '.png',
        'image/gif': '.gif',
        'image/webp': '.webp',
        'image/bmp': '.bmp',
        # Audio (for multimodal models)
        'audio/wav': '.wav',
        'audio/mp3': '.mp3',
        'audio/ogg': '.ogg',
        # Video (for multimodal models)
        'video/mp4': '.mp4',
        'video/mpeg': '.mpeg',
        'video/mov': '.mov',
        'video/avi': '.avi',
    }
    
    # No static models - all models are fetched dynamically from Google API
    
    def __init__(self, api_key: str = None):
        api_key = api_key or os.getenv('GOOGLE_API_KEY')
        if not api_key:
            raise ValueError("Google API key not provided. Set GOOGLE_API_KEY environment variable.")
        super().__init__(api_key)
        
        # Configure the Google AI SDK
        genai.configure(api_key=api_key)
        
    def get_available_models(self) -> List[str]:
        """Get list of available Gemini models from Google API"""
        try:
            dynamic_models = self.get_available_models_from_api()
            if dynamic_models:
                # Extract model names from API response
                model_names = [model['base_model_id'] for model in dynamic_models if model['base_model_id']]
                self.logger.info(f"Fetched {len(model_names)} dynamic Gemini models from API")
                return model_names
            else:
                self.logger.warning("No models returned from Google API")
                return []
        except Exception as e:
            self.logger.error(f"Failed to get dynamic models from Google API: {e}")
            return []
    
    def get_available_models_from_api(self) -> List[dict]:
        """Get list of available Gemini models from Google AI API (no caching - handled by app)"""
        
        try:
            # Use the Google AI API to fetch available models
            import google.generativeai as genai
            
            # Configure the API
            genai.configure(api_key=self.api_key)
            
            # Fetch models from the API
            models_response = genai.list_models()
            
            # Filter for models that support generateContent
            available_models = []
            for model in models_response:
                if hasattr(model, 'supported_generation_methods') and 'generateContent' in model.supported_generation_methods:
                    self.logger.debug(f"Model name: '{model.name}', base_model_id: '{getattr(model, 'base_model_id', 'N/A')}'")
                    
                    # Extract base model ID more robustly
                    base_model_id = getattr(model, 'base_model_id', None)
                    if not base_model_id and model.name:
                        base_model_id = model.name.replace('models/', '')
                    
                    if not base_model_id:
                        self.logger.warning(f"Skipping model with empty name: {model}")
                        continue
                    
                    supported_methods = getattr(model, 'supported_generation_methods', [])
                    supported_modalities = getattr(model, 'supported_modalities', [])
                    self.logger.debug(f"Model {base_model_id} - Methods: {supported_methods}, Modalities: {supported_modalities}")
                    
                    # Determine multimodal capabilities from API response
                    supports_file_upload = self._determine_multimodal_capability(model, base_model_id)
                    
                    model_info = {
                        'name': model.name,
                        'base_model_id': base_model_id,
                        'display_name': getattr(model, 'display_name', base_model_id.replace('-', ' ').title()),
                        'description': getattr(model, 'description', ''),
                        'input_token_limit': getattr(model, 'input_token_limit', 0),
                        'output_token_limit': getattr(model, 'output_token_limit', 8192),
                        'supports_file_upload': supports_file_upload,
                        'thinking': getattr(model, 'thinking', False)
                    }
                    available_models.append(model_info)
            
            # Sort by display name for better UX
            available_models.sort(key=lambda x: x['display_name'])
            
            self.logger.info(f"Fetched {len(available_models)} Gemini models from API")
            return available_models
            
        except Exception as e:
            self.logger.error(f"Error fetching models from API: {e}")
            return []
    
    def validate_file(self, file_path: str) -> bool:
        """Validate if file is supported by Gemini"""
        if not os.path.exists(file_path):
            return False
            
        # Check file size (Gemini API limits: 50MB max, but 20MB recommended for direct upload)
        file_size = os.path.getsize(file_path)
        max_size = 50 * 1024 * 1024  # 50MB limit per API docs
        if file_size > max_size:
            self.logger.warning(f"File {file_path} is {file_size/1024/1024:.1f}MB, exceeds 50MB limit")
            return False
            
        # Note: PDF page count validation removed - Gemini API handles this during upload
            
        mime_type, _ = mimetypes.guess_type(file_path)
        file_extension = Path(file_path).suffix.lower()
        
        # Check if it's a DOCX file (handled via text extraction)
        if self._is_docx_file(file_path):
            return True
        
        return (mime_type in self.SUPPORTED_FILE_TYPES or 
                file_extension in self.SUPPORTED_FILE_TYPES.values())
    
    def _determine_multimodal_capability(self, model, base_model_id: str) -> bool:
        """Determine if a model supports multimodal/file upload capabilities from API response"""
        try:
            # Check for supported generation methods
            supported_methods = getattr(model, 'supported_generation_methods', [])
            if 'generateContent' not in supported_methods:
                return False
            
            # Check for supported modalities (if available)
            supported_modalities = getattr(model, 'supported_modalities', [])
            if supported_modalities:
                # If multiple modalities are supported, it's multimodal
                return len(supported_modalities) > 1
            
            # Check for specific capabilities that indicate multimodal support
            # Look for models that support images, audio, video, or documents
            model_name_lower = base_model_id.lower()
            
            # Known multimodal patterns
            multimodal_patterns = [
                'pro', 'flash', 'vision', 'multimodal', 'image', 'audio', 'video'
            ]
            
            # Exclude text-only models
            text_only_patterns = ['text', 'chat', 'instruct']
            
            # Check if it matches multimodal patterns but not text-only
            has_multimodal_pattern = any(pattern in model_name_lower for pattern in multimodal_patterns)
            has_text_only_pattern = any(pattern in model_name_lower for pattern in text_only_patterns)
            
            # If it has multimodal patterns and doesn't have text-only patterns, it's likely multimodal
            if has_multimodal_pattern and not has_text_only_pattern:
                return True
            
            # If we reach here, we couldn't determine multimodal capability from API
            # Return False to be conservative - only allow file uploads if explicitly detected
            return False
            
        except Exception as e:
            self.logger.debug(f"Error determining multimodal capability for {base_model_id}: {e}")
            # No fallback - rely entirely on API response
            return False
    
    def _supports_file_upload(self, model: str) -> bool:
        """Check if model supports file uploads based on Google API response"""
        try:
            dynamic_models = self.get_available_models_from_api()
            for model_info in dynamic_models:
                if model_info['base_model_id'] == model:
                    return model_info.get('supports_file_upload', False)
        except Exception as e:
            self.logger.debug(f"Error checking file upload support for {model}: {e}")
        
        # If we can't determine from API, assume it doesn't support file uploads
        return False
    
    def _get_model_description(self, model_name: str) -> str:
        """Get description for a model from Google API response"""
        try:
            dynamic_models = self.get_available_models_from_api()
            for model_info in dynamic_models:
                if model_info['base_model_id'] == model_name:
                    return model_info.get('description', f'{model_name.replace("-", " ").title()} model')
        except Exception:
            pass
        
        # Fallback to generic description
        return f'{model_name.replace("-", " ").title()} model'
    
    def _get_model_token_limit(self, model_name: str) -> int:
        """Get token limit for a model from Google API response"""
        try:
            dynamic_models = self.get_available_models_from_api()
            for model_info in dynamic_models:
                if model_info['base_model_id'] == model_name:
                    return model_info.get('input_token_limit', 8192)
        except Exception:
            pass
        
        # Fallback to default token limit
        return 8192
    
    def _is_image_file(self, file_path: str) -> bool:
        """Check if file is an image"""
        mime_type, _ = mimetypes.guess_type(file_path)
        return mime_type and mime_type.startswith('image/')
    
    def _is_audio_file(self, file_path: str) -> bool:
        """Check if file is audio"""
        mime_type, _ = mimetypes.guess_type(file_path)
        return mime_type and mime_type.startswith('audio/')
    
    def _is_video_file(self, file_path: str) -> bool:
        """Check if file is video"""
        mime_type, _ = mimetypes.guess_type(file_path)
        return mime_type and mime_type.startswith('video/')
    
    def _is_docx_file(self, file_path: str) -> bool:
        """Check if file is DOCX"""
        mime_type, _ = mimetypes.guess_type(file_path)
        file_extension = Path(file_path).suffix.lower()
        return (mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or 
                file_extension == '.docx')
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract plain text from DOCX file including headers, body, and footers"""
        try:
            if not DOCX_SUPPORT:
                return f"[DOCX file: {os.path.basename(file_path)}]\nNote: python-docx library not available for text extraction."
            
            self.logger.info(f"Extracting text from DOCX file {file_path}")
            doc = Document(file_path)
            
            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            # Extract headers and footers
            for section in doc.sections:
                # Headers
                if section.header:
                    for paragraph in section.header.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            text_parts.append(f"[Header] {text}")
                
                # Footers
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            text_parts.append(f"[Footer] {text}")
            
            extracted_text = "\n".join(text_parts)
            self.logger.debug(f"Extracted {len(extracted_text)} characters from DOCX")
            
            if not extracted_text.strip():
                return f"[DOCX file: {os.path.basename(file_path)}]\nNote: No text content found in the document."
            
            return extracted_text
            
        except Exception as e:
            self.logger.error(f"Failed to extract text from DOCX file {file_path}: {type(e).__name__}: {e}")
            return f"[DOCX file: {os.path.basename(file_path)}]\nNote: Failed to extract text content: {str(e)}"
    
    
    @BaseProvider._measure_time
    def evaluate_with_file(self, 
                          file_path: str, 
                          prompt: str, 
                          model: str) -> EvaluationResult:
        """Evaluate file with Gemini"""
        if not self.validate_file(file_path):
            return EvaluationResult(
                provider="Gemini",
                model=model,
                output_text="",
                status="error",
                error_message=f"Unsupported file type or file too large: {file_path}"
            )
        
        try:
            # Create the model instance
            generation_config = {
                "temperature": 0.7,
                "top_p": 0.95,
                "top_k": 64,
                "max_output_tokens": 8192,
            }
            
            # Safety settings - allow most content for evaluation purposes
            safety_settings = [
                {
                    "category": HarmCategory.HARM_CATEGORY_HARASSMENT,
                    "threshold": HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE
                },
                {
                    "category": HarmCategory.HARM_CATEGORY_HATE_SPEECH,
                    "threshold": HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE
                },
                {
                    "category": HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT,
                    "threshold": HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE
                },
                {
                    "category": HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT,
                    "threshold": HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE
                },
            ]
            
            gemini_model = genai.GenerativeModel(
                model_name=model,
                generation_config=generation_config,
                safety_settings=safety_settings
            )
            
            # Handle different file types
            if self._is_docx_file(file_path):
                # Handle DOCX files by extracting text content (Gemini doesn't support DOCX uploads)
                try:
                    docx_text = self._extract_docx_text(file_path)
                    content = f"{prompt}\n\nDOCX file content:\n{docx_text}"
                    self.logger.info(f"Processing DOCX file as text: {len(docx_text)} characters extracted")
                    
                except Exception as docx_error:
                    self.logger.error(f"Failed to process DOCX file: {type(docx_error).__name__}: {docx_error}")
                    return EvaluationResult(
                        provider="Gemini",
                        model=model,
                        output_text="",
                        status="error",
                        error_message=f"DOCX file {os.path.basename(file_path)} could not be processed: {str(docx_error)}"
                    )
                    
            elif self._supports_file_upload(model):
                # For multimodal models, upload the file
                self.logger.info(f"Uploading file {file_path} to Gemini...")
                
                try:
                    # Determine MIME type
                    mime_type, _ = mimetypes.guess_type(file_path)
                    if not mime_type:
                        # Set default MIME types based on file extension
                        if file_path.lower().endswith('.pdf'):
                            mime_type = 'application/pdf'
                        elif file_path.lower().endswith('.txt'):
                            mime_type = 'text/plain'
                        else:
                            mime_type = 'application/octet-stream'
                    
                    self.logger.debug(f"Uploading file with MIME type: {mime_type}")
                    
                    # Upload file to Gemini with explicit MIME type
                    uploaded_file = genai.upload_file(file_path, mime_type=mime_type)
                    
                    # Wait for file processing (required for some file types)
                    max_wait_time = 60  # Maximum 60 seconds wait
                    wait_time = 0
                    while uploaded_file.state.name == "PROCESSING" and wait_time < max_wait_time:
                        self.logger.debug(f"File processing... ({wait_time}s)")
                        time.sleep(2)
                        wait_time += 2
                        uploaded_file = genai.get_file(uploaded_file.name)
                    
                    if uploaded_file.state.name == "FAILED":
                        self.logger.error(f"File upload failed: {uploaded_file.state.name}")
                        return EvaluationResult(
                            provider="Gemini",
                            model=model,
                            output_text="",
                            status="error",
                            error_message=f"File upload failed for {file_path}. File state: {uploaded_file.state.name}"
                        )
                    elif uploaded_file.state.name == "PROCESSING":
                        self.logger.warning("File processing timeout")
                        return EvaluationResult(
                            provider="Gemini",
                            model=model,
                            output_text="",
                            status="error",
                            error_message=f"File processing timeout for {file_path}. File state: {uploaded_file.state.name}"
                        )
                    else:
                        # Upload successful, create content with uploaded file
                        self.logger.info(f"File uploaded successfully: {uploaded_file.name}")
                        content = [prompt, uploaded_file]
                    
                except Exception as upload_error:
                    self.logger.error(f"File upload failed: {upload_error}")
                    return EvaluationResult(
                        provider="Gemini",
                        model=model,
                        output_text="",
                        status="error",
                        error_message=f"File upload failed for {file_path}: {str(upload_error)}"
                    )
                        
            else:
                # For text-only models, return error - use multimodal models for file uploads
                return EvaluationResult(
                    provider="Gemini",
                    model=model,
                    output_text="",
                    status="error",
                    error_message=f"Model {model} does not support file uploads. Use a multimodal model like gemini-1.5-pro or gemini-1.5-flash for file processing."
                )
            
            # Generate response
            self.logger.info(f"Generating response with Gemini {model}...")
            self.logger.debug(f"Content type: {type(content)}, Content preview: {str(content)[:200]}...")
            response = gemini_model.generate_content(content)
            
            # Handle response
            if response.candidates and len(response.candidates) > 0:
                candidate = response.candidates[0]
                
                # Check if response was blocked
                if hasattr(candidate, 'finish_reason') and candidate.finish_reason.name in ['SAFETY', 'RECITATION']:
                    return EvaluationResult(
                        provider="Gemini",
                        model=model,
                        output_text="",
                        status="error",
                        error_message=f"Response blocked due to {candidate.finish_reason.name}"
                    )
                
                output_text = candidate.content.parts[0].text if candidate.content.parts else ""
                
                # Extract token usage if available
                input_tokens = None
                output_tokens = None
                if hasattr(response, 'usage_metadata'):
                    input_tokens = getattr(response.usage_metadata, 'prompt_token_count', None)
                    output_tokens = getattr(response.usage_metadata, 'candidates_token_count', None)
                
                return EvaluationResult(
                    provider="Gemini",
                    model=model,
                    output_text=output_text,
                    input_tokens=input_tokens,
                    output_tokens=output_tokens,
                    status="success"
                )
            else:
                return EvaluationResult(
                    provider="Gemini",
                    model=model,
                    output_text="",
                    status="error",
                    error_message="No response generated"
                )
                
        except Exception as e:
            error_msg = str(e)
            self.logger.error(f"Gemini evaluation failed: {error_msg}")
            
            # Handle specific Gemini errors
            if "quota exceeded" in error_msg.lower():
                error_msg = "Gemini API quota exceeded. Check your Google Cloud billing or rate limits."
            elif "invalid api key" in error_msg.lower():
                error_msg = "Invalid Google API key. Check your GOOGLE_API_KEY environment variable."
            elif "permission denied" in error_msg.lower():
                error_msg = "Permission denied. Ensure your API key has access to Gemini models."
            elif "file too large" in error_msg.lower():
                error_msg = f"File {file_path} is too large for Gemini API."
            
            return EvaluationResult(
                provider="Gemini",
                model=model,
                output_text="",
                status="error",
                error_message=error_msg
            )
        
        finally:
            try:
                if 'uploaded_file' in locals() and uploaded_file:
                    genai.delete_file(uploaded_file.name)
                    self.logger.debug(f"Cleaned up uploaded file: {uploaded_file.name}")
            except Exception as cleanup_error:
                self.logger.warning(f"Failed to clean up uploaded file: {cleanup_error}")