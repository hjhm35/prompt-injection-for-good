from typing import List
import os
import base64
import mimetypes
from pathlib import Path

try:
    import anthropic
except ImportError:
    raise ImportError("Anthropic library not installed. Run: pip install anthropic")

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

from .base import BaseProvider, EvaluationResult

class AnthropicProvider(BaseProvider):
    """Anthropic Claude provider implementation"""
    
    SUPPORTED_FILE_TYPES = {
        'application/pdf': '.pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
        'application/msword': '.doc',
        'text/plain': '.txt',
        'text/markdown': '.md',
        'message/rfc822': '.eml',
        'text/email': '.eml',
        'image/jpeg': '.jpg',
        'image/png': '.png',
        'image/gif': '.gif',
        'image/webp': '.webp'
    }
    
    # No static models - all models fetched dynamically from API
    
    def __init__(self, api_key: str = None):
        api_key = api_key or os.getenv('ANTHROPIC_API_KEY')
        if not api_key:
            raise ValueError("Anthropic API key not provided")
        super().__init__(api_key)
        self.client = anthropic.Anthropic(
            api_key=api_key,
            default_headers={"anthropic-beta": "pdfs-2024-09-25"}
        )
        
    def get_available_models(self) -> List[str]:
        """Get list of available Anthropic models (dynamic from API)"""
        try:
            models = self.get_available_models_from_api()
            return [model['id'] for model in models if model.get('id')]
        except Exception as e:
            self.logger.error(f"Failed to fetch Anthropic models from API: {e}")
            return []
    
    def get_available_models_from_api(self) -> List[dict]:
        """Get list of available Anthropic models from Anthropic API (no caching - handled by app)"""
        try:
            self.logger.info("Fetching Anthropic models from API...")
            
            # Call Anthropic API to get models list
            response = self.client.models.list()
            
            available_models = []
            for model in response.data:
                model_info = {
                    'id': model.id,
                    'display_name': model.display_name,
                    'created_at': model.created_at,
                    'type': model.type,
                    'supports_file_upload': self._determine_file_upload_capability_from_api(model),
                    'supports_vision': self._determine_vision_capability_from_api(model),
                    'context_length': self._get_context_length_from_api(model)
                }
                available_models.append(model_info)
            
            self.logger.info(f"Fetched {len(available_models)} Anthropic models from API")
            return available_models
            
        except Exception as e:
            self.logger.error(f"Error fetching Anthropic models from API: {e}")
            raise
    
    def _determine_file_upload_capability_from_api(self, model) -> bool:
        """Determine if model supports file uploads based on API response"""
        # Anthropic models generally support file uploads (PDFs, images, etc.)
        # Based on the model ID patterns and known capabilities
        model_id = model.id.lower()
        
        # All current Anthropic models support file uploads
        # This includes PDFs, images, and other document types
        return True
    
    def _determine_vision_capability_from_api(self, model) -> bool:
        """Determine if model supports vision based on API response"""
        # Anthropic models generally support vision (image processing)
        model_id = model.id.lower()
        
        # All current Anthropic models support vision
        return True
    
    def _get_context_length_from_api(self, model) -> int:
        """Get context length from API response"""
        model_id = model.id.lower()
        
        # Context length mapping based on known Anthropic model capabilities
        if 'claude-3-5-sonnet' in model_id:
            return 200000  # 200K tokens
        elif 'claude-3-5-haiku' in model_id:
            return 200000  # 200K tokens
        elif 'claude-sonnet-4' in model_id:
            return 200000  # 200K tokens
        elif 'claude-opus-4' in model_id:
            return 200000  # 200K tokens
        elif 'claude-3-7-sonnet' in model_id:
            return 200000  # 200K tokens
        else:
            return 200000  # Default to 200K tokens for Anthropic models
    
    def validate_file(self, file_path: str) -> bool:
        """Validate if file is supported by Anthropic"""
        if not os.path.exists(file_path):
            return False
            
        mime_type, _ = mimetypes.guess_type(file_path)
        file_extension = Path(file_path).suffix.lower()
        
        return (mime_type in self.SUPPORTED_FILE_TYPES or 
                file_extension in self.SUPPORTED_FILE_TYPES.values())
    
    def _encode_pdf(self, file_path: str) -> str:
        """Encode PDF file to base64 for direct API submission"""
        try:
            # Check file size first (32MB limit for PDFs)
            file_size = os.path.getsize(file_path)
            self.logger.info(f"Encoding PDF file {file_path} (size: {file_size/1024/1024:.2f}MB)")
            
            if file_size > 32 * 1024 * 1024:  # 32MB limit for PDFs
                raise ValueError(f"PDF file size {file_size/1024/1024:.2f}MB exceeds 32MB limit")
            
            with open(file_path, 'rb') as file:
                file_data = file.read()
                encoded_data = base64.b64encode(file_data).decode('utf-8')
            
            self.logger.info(f"Successfully encoded PDF file to base64 ({len(encoded_data)} characters)")
            return encoded_data
        except Exception as e:
            self.logger.error(f"Failed to encode PDF file {file_path}: {type(e).__name__}: {e}")
            raise
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract plain text from DOCX file including headers, body, and footers"""
        try:
            if not DOCX_SUPPORT:
                return f"[DOCX file: {os.path.basename(file_path)}]\nNote: python-docx library not available for text extraction."
            
            self.logger.info(f"Extracting text from DOCX file {file_path}")
            
            # Open and read the DOCX file
            doc = Document(file_path)
            
            text_content = []
            
            # Extract text from headers and footers in all sections
            for section_idx, section in enumerate(doc.sections):
                # Extract header content
                if section.header:
                    header_text = []
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            header_text.append(paragraph.text.strip())
                    
                    # Extract text from tables in headers
                    for table in section.header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    header_text.append(cell.text.strip())
                    
                    if header_text:
                        text_content.append(f"=== HEADER SECTION {section_idx + 1} ===")
                        text_content.extend(header_text)
                
                # Extract footer content
                if section.footer:
                    footer_text = []
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            footer_text.append(paragraph.text.strip())
                    
                    # Extract text from tables in footers
                    for table in section.footer.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    footer_text.append(cell.text.strip())
                    
                    if footer_text:
                        text_content.append(f"=== FOOTER SECTION {section_idx + 1} ===")
                        text_content.extend(footer_text)
            
            # Extract text from main document body paragraphs
            body_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Only add non-empty paragraphs
                    body_text.append(paragraph.text.strip())
            
            if body_text:
                text_content.append("=== DOCUMENT BODY ===")
                text_content.extend(body_text)
            
            # Extract text from tables in document body
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            table_text.append(cell.text.strip())
            
            if table_text:
                text_content.append("=== DOCUMENT TABLES ===")
                text_content.extend(table_text)
            
            extracted_text = '\n\n'.join(text_content)
            
            self.logger.info(f"Successfully extracted {len(extracted_text)} characters from DOCX file (including headers/footers)")
            return extracted_text
            
        except Exception as e:
            self.logger.error(f"Failed to extract text from DOCX file {file_path}: {type(e).__name__}: {e}")
            return f"[DOCX file: {os.path.basename(file_path)}]\nNote: Could not extract text content - {str(e)}"
    
    @BaseProvider._measure_time
    def evaluate_with_file(self, 
                          file_path: str, 
                          prompt: str, 
                          model: str) -> EvaluationResult:
        """Evaluate file with Anthropic Claude using Files API"""
        if not self.validate_file(file_path):
            return EvaluationResult(
                provider="Anthropic",
                model=model,
                output_text="",
                status="error",
                error_message=f"Unsupported file type: {file_path}"
            )
        
        try:
            mime_type, _ = mimetypes.guess_type(file_path)
            file_extension = Path(file_path).suffix.lower()
            
            # Determine if file is an image
            is_image = mime_type and mime_type.startswith('image/')
            
            # Determine if file is a PDF
            is_pdf = mime_type == 'application/pdf' or file_extension == '.pdf'
            
            # Determine if file is a DOCX
            is_docx = (mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or 
                      file_extension == '.docx')
            
            if is_image:
                # Handle image files with base64 encoding (existing method)
                with open(file_path, 'rb') as file:
                    file_data = file.read()
                    encoded_data = base64.b64encode(file_data).decode('utf-8')
                
                messages = [
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image",
                                "source": {
                                    "type": "base64",
                                    "media_type": mime_type,
                                    "data": encoded_data
                                }
                            },
                            {
                                "type": "text",
                                "text": prompt
                            }
                        ]
                    }
                ]
            elif is_pdf:
                # Handle PDF files using base64 encoding and document content blocks
                try:
                    encoded_data = self._encode_pdf(file_path)
                    
                    messages = [
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "document",
                                    "source": {
                                        "type": "base64",
                                        "media_type": "application/pdf",
                                        "data": encoded_data
                                    }
                                },
                                {
                                    "type": "text",
                                    "text": prompt
                                }
                            ]
                        }
                    ]
                except Exception as encode_error:
                    self.logger.error(f"Failed to encode PDF file: {type(encode_error).__name__}: {encode_error}")
                    # Fallback to text description with specific error
                    error_msg = str(encode_error)
                    if "32MB" in error_msg:
                        fallback_msg = f"PDF file {os.path.basename(file_path)} is too large (exceeds 32MB limit)."
                    else:
                        fallback_msg = f"PDF file {os.path.basename(file_path)} could not be processed: {error_msg}"
                    
                    messages = [
                        {
                            "role": "user",
                            "content": f"{prompt}\n\nNote: {fallback_msg}"
                        }
                    ]
            elif is_docx:
                # Handle DOCX files by extracting text content (Claude API requirement)
                try:
                    docx_text = self._extract_docx_text(file_path)
                    
                    messages = [
                        {
                            "role": "user",
                            "content": f"{prompt}\n\nDOCX file content:\n{docx_text}"
                        }
                    ]
                except Exception as docx_error:
                    self.logger.error(f"Failed to process DOCX file: {type(docx_error).__name__}: {docx_error}")
                    messages = [
                        {
                            "role": "user",
                            "content": f"{prompt}\n\nNote: DOCX file {os.path.basename(file_path)} could not be processed: {str(docx_error)}"
                        }
                    ]
            else:
                # Handle other document files (TXT, EMAIL, etc.) by converting to text
                try:
                    if file_path.endswith(('.txt', '.md', '.eml')):
                        with open(file_path, 'r', encoding='utf-8') as f:
                            file_content = f.read()
                    else:
                        # For other file types, mention the file
                        file_content = f"[File: {os.path.basename(file_path)} - {mime_type or 'unknown type'}]"
                        
                    messages = [
                        {
                            "role": "user",
                            "content": f"{prompt}\n\nFile content:\n{file_content[:10000]}"  # Limit content
                        }
                    ]
                except Exception:
                    messages = [
                        {
                            "role": "user", 
                            "content": f"{prompt}\n\nNote: File {os.path.basename(file_path)} could not be read directly."
                        }
                    ]
            
            # Make API call
            response = self.client.messages.create(
                model=model,
                max_tokens=4000,
                messages=messages
            )
            
            # Extract response data
            output_text = response.content[0].text if response.content else ""
            
            # Get token usage if available
            input_tokens = response.usage.input_tokens if hasattr(response, 'usage') else None
            output_tokens = response.usage.output_tokens if hasattr(response, 'usage') else None
            
            return EvaluationResult(
                provider="Anthropic",
                model=model,
                output_text=output_text,
                input_tokens=input_tokens,
                output_tokens=output_tokens,
                status="success"
            )
            
        except Exception as e:
            self.logger.error(f"Anthropic evaluation failed: {e}")
            return EvaluationResult(
                provider="Anthropic", 
                model=model,
                output_text="",
                status="error",
                error_message=str(e)
            )