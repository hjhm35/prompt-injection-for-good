from typing import List
import os
import base64
import mimetypes
from pathlib import Path
import requests

try:
    from openai import OpenAI
except ImportError:
    raise ImportError("OpenAI library not installed. Run: pip install openai")

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

from .base import BaseProvider, EvaluationResult

class TogetherProvider(BaseProvider):
    """Together AI provider implementation supporting Llama, Qwen, Mistral, Kimi and other models"""
    
    # Together AI supports various file types for vision and document processing
    SUPPORTED_FILE_TYPES = {
        # Documents
        'application/pdf': '.pdf',
        'text/plain': '.txt',
        'text/markdown': '.md',
        'message/rfc822': '.eml',
        'text/email': '.eml',
        'application/json': '.json',
        'text/csv': '.csv',
        # Images  
        'image/jpeg': '.jpg',
        'image/png': '.png',
        'image/gif': '.gif',
        'image/webp': '.webp',
        'image/bmp': '.bmp',
    }
    
    # No static models - all models fetched dynamically from Together.ai API
    # Supports Llama, Qwen, Mistral, Kimi and other model families
    
    def __init__(self, api_key: str = None):
        api_key = api_key or os.getenv('TOGETHER_API_KEY')
        if not api_key:
            raise ValueError("Together AI API key not provided. Set TOGETHER_API_KEY environment variable.")
        super().__init__(api_key)
        
        # Use Together AI's OpenAI-compatible endpoint
        self.client = OpenAI(
            api_key=api_key,
            base_url="https://api.together.xyz/v1"
        )
        
        # For direct API calls, store base URL
        self.base_url = "https://api.together.xyz/v1"
        
    def get_available_models(self) -> List[str]:
        """Get list of available models from Together AI (filtered for mistralai, meta-llama, Qwen, and moonshotai families)"""
        try:
            models = self.get_available_models_from_api()
            
            # Filter for specific model families: mistralai, meta-llama, Qwen, moonshotai
            filtered_models = []
            for model in models:
                if model.get('id'):
                    model_id = model['id']
                    if any(family in model_id for family in ['mistralai', 'meta-llama', 'Qwen', 'moonshotai']):
                        filtered_models.append(model_id)
            
            # Sort the filtered list alphabetically
            filtered_models.sort()
            
            self.logger.info(f"Filtered and sorted {len(filtered_models)} models from {len(models)} total models")
            return filtered_models
            
        except Exception as e:
            self.logger.error(f"Failed to fetch models from Together AI API: {e}")
            return []
    
    def get_available_models_from_api(self) -> List[dict]:
        """Get list of available models from Together.ai API (no caching - handled by app)"""
        try:
            self.logger.info("Fetching models from Together.ai API...")
            
            # Call Together.ai API to get models list using requests directly
            # The OpenAI client has issues with Together AI's response format
            response = requests.get(
                f"{self.base_url}/models",
                headers={"Authorization": f"Bearer {self.api_key}"},
                timeout=30
            )
            
            if response.status_code != 200:
                raise Exception(f"API request failed: {response.status_code} {response.text}")
            
            models_json = response.json()
            
            available_models = []
            
            # Handle response - Together AI can return data as a list or with a data field
            models_data = models_json.get('data', models_json) if isinstance(models_json, dict) else models_json
            
            for model in models_data:
                try:
                    # Handle both object and dict formats
                    if hasattr(model, 'id'):
                        model_id = model.id
                        model_dict = {
                            'id': model.id,
                            'object': getattr(model, 'object', 'model'),
                            'created': getattr(model, 'created', None),
                            'type': getattr(model, 'type', ''),
                            'display_name': getattr(model, 'display_name', model.id),
                            'organization': getattr(model, 'organization', ''),
                            'context_length': getattr(model, 'context_length', 0),
                            'pricing': getattr(model, 'pricing', None)
                        }
                    else:
                        # Handle dict format
                        model_id = model.get('id', str(model))
                        model_dict = {
                            'id': model_id,
                            'object': model.get('object', 'model'),
                            'created': model.get('created', None),
                            'type': model.get('type', ''),
                            'display_name': model.get('display_name', model_id),
                            'organization': model.get('organization', ''),
                            'context_length': model.get('context_length', 0),
                            'pricing': model.get('pricing', None)
                        }
                    
                    # Add computed fields
                    model_dict.update({
                        'supports_file_upload': self._determine_file_upload_capability_from_api(model_dict),
                        'supports_vision': self._determine_vision_capability_from_api(model_dict),
                        'model_family': self._determine_model_family(model_id)
                    })
                    
                    available_models.append(model_dict)
                    
                except Exception as e:
                    self.logger.debug(f"Error processing model {model}: {e}")
                    continue
            
            self.logger.info(f"Fetched {len(available_models)} models from Together.ai API")
            return available_models
            
        except Exception as e:
            self.logger.error(f"Error fetching models from Together.ai API: {e}")
            raise
    
    def _determine_model_family(self, model_id: str) -> str:
        """Determine the model family from model ID"""
        model_id_lower = model_id.lower()
        
        if 'llama' in model_id_lower:
            return 'llama'
        elif 'qwen' in model_id_lower:
            return 'qwen'  
        elif 'mistral' in model_id_lower or 'mixtral' in model_id_lower:
            return 'mistral'
        elif 'kimi' in model_id_lower or 'moonshot' in model_id_lower or 'moonshotai' in model_id_lower:
            return 'kimi'
        elif 'yi' in model_id_lower:
            return 'yi'
        else:
            return 'other'
    
    def _determine_file_upload_capability_from_api(self, model) -> bool:
        """Determine if model supports file uploads based on API response"""
        # Together.ai models generally support file uploads for vision models
        # Check if it's a vision model based on name patterns
        model_id = (model.get('id') if isinstance(model, dict) else getattr(model, 'id', '')).lower()
        model_type = (model.get('type') if isinstance(model, dict) else getattr(model, 'type', '')).lower()
        
        # Vision models typically support file uploads
        if 'vision' in model_id or 'vl' in model_id or model_type == 'vision':
            return True
        
        # Check for known vision model patterns
        vision_patterns = ['vision', 'multimodal', 'image', '-vl-', 'scout', 'maverick']
        return any(pattern in model_id for pattern in vision_patterns)
    
    def _determine_vision_capability_from_api(self, model) -> bool:
        """Determine if model supports vision based on API response"""
        # Check if it's a vision model based on name patterns
        model_id = (model.get('id') if isinstance(model, dict) else getattr(model, 'id', '')).lower()
        model_type = (model.get('type') if isinstance(model, dict) else getattr(model, 'type', '')).lower()
        
        # Vision models support image processing
        if 'vision' in model_id or 'vl' in model_id or model_type == 'vision':
            return True
        
        # Check for known vision model patterns
        vision_patterns = ['vision', 'multimodal', 'image', '-vl-', 'scout', 'maverick']
        return any(pattern in model_id for pattern in vision_patterns)
    
    def validate_file(self, file_path: str) -> bool:
        """Validate if file is supported by Together AI models"""
        if not os.path.exists(file_path):
            return False
            
        # Check file size (Together AI has file size limits)
        file_size = os.path.getsize(file_path)
        max_size = 100 * 1024 * 1024  # 100MB limit (to be verified)
        if file_size > max_size:
            self.logger.warning(f"File {file_path} is {file_size/1024/1024:.1f}MB, may exceed size limits")
            
        mime_type, _ = mimetypes.guess_type(file_path)
        file_extension = Path(file_path).suffix.lower()
        
        # Check if it's a DOCX file (handled via text extraction)
        if self._is_docx_file(file_path):
            return True
        
        # Check if it's a PDF file (handled via text extraction)
        if self._is_pdf_file(file_path):
            return True
        
        return (mime_type in self.SUPPORTED_FILE_TYPES or 
                file_extension in self.SUPPORTED_FILE_TYPES.values())
    
    def _is_vision_model(self, model: str) -> bool:
        """Check if model supports vision/image processing (dynamic detection)"""
        # Check for known vision model patterns
        model_lower = model.lower()
        vision_patterns = ['vision', 'multimodal', 'image']
        return any(pattern in model_lower for pattern in vision_patterns)
    
    def _is_image_file(self, file_path: str) -> bool:
        """Check if file is an image"""
        mime_type, _ = mimetypes.guess_type(file_path)
        return mime_type and mime_type.startswith('image/')
    
    def _upload_file_to_together(self, file_path: str) -> str:
        """Upload file to Together AI and return file ID"""
        try:
            self.logger.info(f"Uploading file {file_path} to Together AI")
            
            with open(file_path, 'rb') as file:
                files = {
                    'file': (os.path.basename(file_path), file, mimetypes.guess_type(file_path)[0])
                }
                
                # Include the required purpose parameter
                data = {
                    'purpose': 'assistants'  # For use with chat completions/assistants
                }
                
                response = requests.post(
                    f"{self.base_url}/files",
                    headers={"Authorization": f"Bearer {self.api_key}"},
                    files=files,
                    data=data,
                    timeout=60  # Longer timeout for file uploads
                )
                
                if response.status_code != 200:
                    raise Exception(f"File upload failed: {response.status_code} {response.text}")
                
                result = response.json()
                file_id = result.get('id')
                
                if not file_id:
                    raise Exception(f"No file ID returned from upload: {result}")
                
                self.logger.info(f"Successfully uploaded file {file_path} with ID: {file_id}")
                return file_id
                
        except Exception as e:
            self.logger.error(f"Failed to upload file {file_path}: {e}")
            raise
    
    def _supports_file_upload(self, file_path: str) -> bool:
        """Check if file type supports Together AI file upload API"""
        mime_type, _ = mimetypes.guess_type(file_path)
        file_extension = Path(file_path).suffix.lower()
        
        # Together AI file upload API supports these types
        upload_supported_types = {
            # Images
            'image/jpeg', 'image/png', 'image/gif', 'image/webp', 'image/bmp',
            # Documents that Together AI can process directly
            'application/pdf', 'text/plain', 'text/markdown', 'message/rfc822', 'text/email', 'application/json', 'text/csv'
        }
        
        upload_supported_extensions = {
            '.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp',
            '.pdf', '.txt', '.md', '.eml', '.json', '.csv'
        }
        
        return (mime_type in upload_supported_types or 
                file_extension in upload_supported_extensions)
    
    def _encode_image(self, file_path: str) -> str:
        """Encode image file to base64 (fallback method)"""
        with open(file_path, 'rb') as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')
    
    def _is_docx_file(self, file_path: str) -> bool:
        """Check if file is DOCX"""
        mime_type, _ = mimetypes.guess_type(file_path)
        file_extension = Path(file_path).suffix.lower()
        return (mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or 
                file_extension == '.docx')
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract plain text from DOCX file including headers, body, and footers"""
        try:
            if not DOCX_SUPPORT:
                return f"[DOCX file: {os.path.basename(file_path)}]\nNote: python-docx library not available for text extraction."
            
            self.logger.info(f"Extracting text from DOCX file {file_path}")
            doc = Document(file_path)
            
            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            # Extract headers and footers
            for section in doc.sections:
                # Headers
                if section.header:
                    for paragraph in section.header.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            text_parts.append(f"[Header] {text}")
                
                # Footers
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            text_parts.append(f"[Footer] {text}")
            
            extracted_text = "\n".join(text_parts)
            self.logger.debug(f"Extracted {len(extracted_text)} characters from DOCX")
            
            if not extracted_text.strip():
                return f"[DOCX file: {os.path.basename(file_path)}]\nNote: No text content found in the document."
            
            return extracted_text
            
        except Exception as e:
            self.logger.error(f"Failed to extract text from DOCX file {file_path}: {type(e).__name__}: {e}")
            return f"[DOCX file: {os.path.basename(file_path)}]\nNote: Failed to extract text content: {str(e)}"
    
    def _is_pdf_file(self, file_path: str) -> bool:
        """Check if file is PDF"""
        mime_type, _ = mimetypes.guess_type(file_path)
        file_extension = Path(file_path).suffix.lower()
        return (mime_type == 'application/pdf' or file_extension == '.pdf')
    
    def _extract_pdf_text(self, file_path: str, max_chars: int = 50000) -> str:
        """Extract plain text from PDF file with character limit to prevent token overflow"""
        try:
            if not PDF_SUPPORT:
                return f"[PDF file: {os.path.basename(file_path)}]\nNote: PyPDF2 library not available for text extraction."
            
            self.logger.info(f"Extracting text from PDF file {file_path}")
            
            text_parts = []
            total_chars = 0
            
            with open(file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                total_pages = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            # Check if adding this page would exceed limit
                            if total_chars + len(page_text) > max_chars:
                                remaining_chars = max_chars - total_chars
                                if remaining_chars > 100:  # Only add if we have reasonable space left
                                    truncated_text = page_text[:remaining_chars]
                                    text_parts.append(f"[Page {page_num + 1}] {truncated_text}")
                                    text_parts.append(f"\n[TRUNCATED - Reached {max_chars} character limit. PDF has {total_pages} total pages.]")
                                break
                            else:
                                text_parts.append(f"[Page {page_num + 1}] {page_text}")
                                total_chars += len(page_text)
                    except Exception as page_error:
                        self.logger.debug(f"Error extracting page {page_num + 1}: {page_error}")
                        continue
            
            extracted_text = "\n".join(text_parts)
            self.logger.info(f"Extracted {len(extracted_text)} characters from PDF ({len(pdf_reader.pages)} pages)")
            
            if not extracted_text.strip():
                return f"[PDF file: {os.path.basename(file_path)}]\nNote: No text content found in the document."
            
            return extracted_text
            
        except Exception as e:
            self.logger.error(f"Failed to extract text from PDF file {file_path}: {type(e).__name__}: {e}")
            return f"[PDF file: {os.path.basename(file_path)}]\nNote: Failed to extract text content: {str(e)}"
    
    def _read_text_file(self, file_path: str) -> str:
        """Read text content from file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
    
    @BaseProvider._measure_time
    def evaluate_with_file(self, 
                          file_path: str, 
                          prompt: str, 
                          model: str) -> EvaluationResult:
        """Evaluate file with Together AI models (Llama, Qwen, Mistral, Kimi, etc.)"""
        if not self.validate_file(file_path):
            return EvaluationResult(
                provider="Together",
                model=model,
                output_text="",
                status="error",
                error_message=f"Unsupported file type: {file_path}"
            )
        
        try:
            is_vision_model = self._is_vision_model(model)
            is_image = self._is_image_file(file_path)
            
            if is_image and not is_vision_model:
                return EvaluationResult(
                    provider="Together",
                    model=model,
                    output_text="",
                    status="error",
                    error_message=f"Model {model} does not support image files. Use a vision model like Qwen2.5-VL-72B-Instruct or Llama-3.2-11B-Vision-Instruct-Turbo"
                )
            
            # Use content extraction and base64 encoding for all file types
            # Together AI's file upload API is primarily for fine-tuning, not chat completions
            
            # Handle file processing via content extraction and encoding
            if self._is_image_file(file_path):
                if not is_vision_model:
                    return EvaluationResult(
                        provider="Together",
                        model=model,
                        output_text="",
                        status="error",
                        error_message=f"Image file {file_path} requires a vision model"
                    )
                
                # Use base64 encoding for images
                mime_type, _ = mimetypes.guess_type(file_path)
                base64_image = self._encode_image(file_path)
                
                messages = [
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": prompt
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:{mime_type};base64,{base64_image}"
                                }
                            }
                        ]
                    }
                ]
                
            elif self._is_docx_file(file_path):
                # Handle DOCX files by extracting text content
                try:
                    file_content = self._extract_docx_text(file_path)
                    combined_prompt = f"DOCX file content:\n{file_content}\n\nPrompt: {prompt}"
                    
                except Exception as docx_error:
                    return EvaluationResult(
                        provider="Together",
                        model=model,
                        output_text="",
                        status="error",
                        error_message=f"DOCX file {os.path.basename(file_path)} could not be processed: {str(docx_error)}"
                    )
                
                messages = [
                    {
                        "role": "user",
                        "content": combined_prompt
                    }
                ]
                
            elif self._is_pdf_file(file_path):
                # Handle PDF files by extracting text content with token limit protection
                try:
                    file_content = self._extract_pdf_text(file_path)
                    combined_prompt = f"PDF file content:\n{file_content}\n\nPrompt: {prompt}"
                    
                except Exception as pdf_error:
                    return EvaluationResult(
                        provider="Together",
                        model=model,
                        output_text="",
                        status="error",
                        error_message=f"PDF file {os.path.basename(file_path)} could not be processed: {str(pdf_error)}"
                    )
                
                messages = [
                    {
                        "role": "user",
                        "content": combined_prompt
                    }
                ]
                
            else:
                # Read regular text file content and combine with prompt
                file_content = self._read_text_file(file_path)
                combined_prompt = f"File content:\n{file_content}\n\nPrompt: {prompt}"
                
                messages = [
                    {
                        "role": "user",
                        "content": combined_prompt
                    }
                ]
            
            # Make API request
            response = self.client.chat.completions.create(
                model=model,
                messages=messages,
                temperature=0.7,
                max_tokens=4000,
                stream=False
            )
            
            output_text = response.choices[0].message.content
            input_tokens = getattr(response.usage, 'prompt_tokens', None) if hasattr(response, 'usage') else None
            output_tokens = getattr(response.usage, 'completion_tokens', None) if hasattr(response, 'usage') else None
            
            return EvaluationResult(
                provider="Together",
                model=model,
                output_text=output_text,
                input_tokens=input_tokens,
                output_tokens=output_tokens,
                status="success"
            )
            
        except Exception as e:
            error_msg = str(e)
            self.logger.error(f"Together AI evaluation failed: {error_msg}")
            
            # Handle specific Together AI errors
            if "insufficient_quota" in error_msg.lower():
                error_msg = "Together AI quota exceeded. Check your account balance or upgrade plan."
            elif "invalid_request_error" in error_msg.lower():
                error_msg = f"Invalid request to Together AI: {error_msg}"
            elif "model_not_found" in error_msg.lower():
                error_msg = f"Model {model} not found on Together AI. Check available models with 'llm-eval list-models'"
            
            return EvaluationResult(
                provider="Together",
                model=model,
                output_text="",
                status="error",
                error_message=error_msg
            )